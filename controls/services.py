"""
services.py
-----------
Core business logic: importing the IKS Excel control list, parsing
Arbeitspapier (working paper) Word documents, and generating the final
Word report from a docxtpl template.
"""

import re
import io
import json
import openpyxl
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate

from .models import Control, ArbeitspapierSubmission, ChangeLogEntry

SECTION_HEADER_RE = re.compile(r"^(\d+(?:\.\d+)?)\s*\t")

# Some working papers omit the "5.1" numeric prefix on this specific section
# header (just the label, e.g. "       Durchgeführte Testaktivitäten" with no
# leading "5.1\t") -- recognized as a fallback by label text when the numeric
# pattern above doesn't match.
KNOWN_SECTION_LABELS = {
    "durchgeführte testaktivitäten": "5.1",
}

# Any of these substrings (case-insensitive) in section 5.3 means "no finding".
NO_DEVIATION_MARKERS = [
    "keine feststellung",
    "keine abweichung",
    "no deviation",
    "no exception",
]


def _normalize(text):
    """Lowercase + collapse whitespace, for tolerant text comparison."""
    return re.sub(r"\s+", " ", (text or "").strip().lower())


# Human-readable explanation of what each validation check actually verifies --
# shown as an info-icon tooltip next to each result in the Controls table.
CHECK_DESCRIPTIONS = {
    "kontrollziel_match": "Compares the Kontrollziel text in this working paper's section 1.2 against the Kontrollziel already recorded for this control from the IKS master Excel list.",
    "kontrollbeschreibung_match": "Compares the Kontrollbeschreibung text in this working paper's section 1.3 against the Kontrollbeschreibung already recorded for this control from the IKS master Excel list.",
    "standard_phrasing": "Checks whether section 5.1 (test activities) uses standard audit phrasing: \"Discussion with the process owner to...\", \"Inspection of existing process documentation...\", and \"Review of an example of ... with regard to the following aspects:\".",
    "writing_style": "Planned: compares the phrasing of section 5.1 against other reports (via embeddings) to flag inconsistent tone or terminology. Not yet implemented.",
}


# ---------------------------------------------------------------------------
# Validation framework
#
# Each validator takes (control, parsed_dict) and returns one result dict:
#   {"check": <key>, "passed": bool, "message": <human-readable explanation>}
#
# `control` here is the Control row's state BEFORE this import touches it
# (so we're comparing the working paper against whatever the master IKS
# Excel import already established -- if Excel hasn't been imported yet,
# checks are skipped rather than failed).
#
# Add new checks by writing a function with this same signature and adding
# it to VALIDATORS below. This is where a future "writing style" check
# (e.g. comparing 5.1 phrasing against other reports via embeddings) would
# plug in.
# ---------------------------------------------------------------------------

def _validate_kontrollziel_match(control_before, parsed, language):
    master_value = control_before.get("kontrollziel", "")
    desc = CHECK_DESCRIPTIONS["kontrollziel_match"]
    if not master_value:
        return {"check": "kontrollziel_match", "passed": True, "description": desc,
                "message": "No master Kontrollziel to compare (import the IKS Excel first for a real check)."}
    match = _normalize(master_value) == _normalize(parsed["kontrollziel"])
    if match:
        return {"check": "kontrollziel_match", "passed": True, "description": desc, "message": "Kontrollziel matches the IKS master list."}
    return {"check": "kontrollziel_match", "passed": False, "description": desc,
            "message": f'Kontrollziel differs from IKS master list. Master: "{master_value}" — Working paper: "{parsed["kontrollziel"]}"'}


def _validate_kontrollbeschreibung_match(control_before, parsed, language):
    master_value = control_before.get("kontrollbeschreibung", "")
    desc = CHECK_DESCRIPTIONS["kontrollbeschreibung_match"]
    if not master_value:
        return {"check": "kontrollbeschreibung_match", "passed": True, "description": desc,
                "message": "No master Kontrollbeschreibung to compare (import the IKS Excel first for a real check)."}
    match = _normalize(master_value) == _normalize(parsed["kontrollbeschreibung"])
    if match:
        return {"check": "kontrollbeschreibung_match", "passed": True, "description": desc, "message": "Kontrollbeschreibung matches the IKS master list."}
    return {"check": "kontrollbeschreibung_match", "passed": False, "description": desc,
            "message": "Kontrollbeschreibung differs from the IKS master list — review the wording."}


# Roadmap / not yet implemented: writing-style consistency check for section 5.1
# (e.g. via embeddings, comparing phrasing against other reports' test activities
# to flag inconsistent tone/terminology). Left as a placeholder so it's easy to
# wire in later without changing the calling code in import_arbeitspapier().
def _validate_writing_style(control_before, parsed, language):
    return {"check": "writing_style", "passed": True, "description": CHECK_DESCRIPTIONS["writing_style"],
            "message": "Not yet implemented — planned: compare 5.1 phrasing against other reports via embeddings."}


# Standard audit phrasing expected in section 5.1 test activities. Configured
# per-language via the StandardPhrase model (see "Standard Phrases" in the
# sidebar) rather than hardcoded here, since EN/DE wording differs.

def _validate_standard_phrasing(control_before, parsed, language):
    """
    Returns ONE validation result PER configured phrase (not a single
    aggregated pass/fail), so the Controls table shows exactly which
    phrase(s) are missing rather than one vague "standard_phrasing" failure.
    """
    from .models import StandardPhrase
    try:
        activities = json.loads(parsed["test_activities"])
    except (json.JSONDecodeError, TypeError):
        activities = []
    flattened = " ".join(item.get("text", "") for item in activities).lower()

    phrases = list(StandardPhrase.objects.filter(language=language, active=True))
    if not phrases:
        return [{"check": "standard_phrasing", "passed": True,
                  "description": CHECK_DESCRIPTIONS["standard_phrasing"],
                  "message": f"No standard phrases configured for language '{language}' yet -- add some on the Standard Phrases page."}]

    results = []
    for sp in phrases:
        try:
            found = bool(re.search(sp.phrase, flattened, re.IGNORECASE | re.DOTALL))
        except re.error:
            found = False  # malformed pattern -- treat as not found rather than crash the import
        label = sp.phrase if len(sp.phrase) <= 60 else sp.phrase[:57] + "..."
        results.append({
            "check": label,
            "passed": found,
            "description": f'Checks whether section 5.1 test activities contain wording matching: "{sp.phrase}"',
            "message": "Found in test activities." if found else f'Not found — no text in the test activities matches "{sp.phrase}".',
        })
    return results


VALIDATORS = [
    _validate_kontrollziel_match,
    _validate_kontrollbeschreibung_match,
    _validate_standard_phrasing,
    # _validate_writing_style,  # enable once implemented
]


def run_validations(control_before, parsed, language):
    """
    Runs every validator and flattens the results into one flat list --
    most validators return a single result dict, but some (like
    standard phrasing) return a list of per-item results.
    """
    results = []
    for v in VALIDATORS:
        r = v(control_before, parsed, language)
        if isinstance(r, list):
            results.extend(r)
        else:
            results.append(r)
    return results


# ---------------------------------------------------------------------------
# 1. Master IKS control list import (Excel)
# ---------------------------------------------------------------------------

def import_iks_excel(project, file_obj):
    """
    Reads an IKS.xlsx-style file (columns: ID, Kontrollziel, Kontrollbeschreibung)
    and creates/updates Control rows for the given project.
    Returns the number of controls imported.
    """
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active

    header = [str(c.value).strip() if c.value else "" for c in ws[1]]
    try:
        id_col = header.index("ID")
        ziel_col = header.index("Kontrollziel")
        besch_col = header.index("Kontrollbeschreibung")
    except ValueError as e:
        raise ValueError(
            f"Expected columns 'ID', 'Kontrollziel', 'Kontrollbeschreibung' — found {header}"
        ) from e

    count = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        control_id = row[id_col]
        if not control_id:
            continue
        Control.objects.update_or_create(
            project=project,
            control_id=str(control_id).strip(),
            defaults={
                "kontrollziel": (row[ziel_col] or "").strip() if row[ziel_col] else "",
                "kontrollbeschreibung": (row[besch_col] or "").strip() if row[besch_col] else "",
            },
        )
        count += 1
    return count


# ---------------------------------------------------------------------------
# 2. Arbeitspapier (working paper) parsing
# ---------------------------------------------------------------------------

GEPRUEFT_RE = re.compile(r"Gepr[üu]ft:\s*(.+)", re.IGNORECASE)
DATUM_RE = re.compile(r"Datum:\s*(\d{1,2})\.(\d{1,2})\.(\d{4})")


def _extract_geprueft_metadata(doc):
    """
    Reads the working paper's metadata table (e.g. "Geprüft: FNO" / "Datum:
    31.07.2025") and returns (geprueft_von, geprueft_datum) -- the auditor's
    initials/name and the review date, straight from the working paper
    itself rather than a manual "mark reviewed" click in the app.
    """
    from datetime import date
    geprueft_von = ""
    geprueft_datum = None

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                m = GEPRUEFT_RE.match(text)
                if m:
                    geprueft_von = m.group(1).strip()
                m2 = DATUM_RE.search(text)
                if m2:
                    day, month, year = m2.groups()
                    try:
                        geprueft_datum = date(int(year), int(month), int(day))
                    except ValueError:
                        pass
    return geprueft_von, geprueft_datum


def _extract_sections(doc):
    """
    Walks every paragraph in the document. A paragraph that starts with a
    "N.M<TAB>" or "N.<TAB>" pattern (e.g. "1.1\tKontrollnummer...", "5.\t...")
    marks the start of a new numbered section. Everything after it (until the
    next such marker) is that section's content.

    Returns: {section_number: [(text, ilvl, left_indent, num_id), ...]}
    ilvl is the Word list indentation level (0 = top-level bullet,
    1 = sub-bullet, None = not a list item at all -- e.g. plain paragraphs).
    left_indent and num_id are captured too, since some documents encode
    sub-bullets differently (see _resolve_bullet_levels below).
    """
    sections = {}
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        match = SECTION_HEADER_RE.match(text)
        if match:
            current_section = match.group(1)
            sections[current_section] = []
        elif text.lower() in KNOWN_SECTION_LABELS and KNOWN_SECTION_LABELS[text.lower()] not in sections:
            current_section = KNOWN_SECTION_LABELS[text.lower()]
            sections[current_section] = []
        elif current_section:
            ilvl = None
            left_indent = None
            num_id = None
            pPr = para._p.pPr
            if pPr is not None:
                if pPr.numPr is not None:
                    if pPr.numPr.ilvl is not None:
                        ilvl = pPr.numPr.ilvl.val
                    if pPr.numPr.numId is not None:
                        num_id = pPr.numPr.numId.val
                if pPr.ind is not None and pPr.ind.left is not None:
                    left_indent = pPr.ind.left
            sections[current_section].append((text, ilvl, left_indent, num_id))

    return sections


def _resolve_bullet_levels(items):
    """
    Given a section's raw (text, ilvl, left_indent, num_id) tuples, resolves
    each line to a final nesting level (0 = main bullet, 1 = sub-bullet).

    Real working papers encode this three different ways, checked in order:
      1. Word's proper multilevel list ilvl (ilvl > 0 means sub-bullet) --
         the standard case.
      2. Plain paragraph indentation, when ilvl stays 0 throughout but some
         lines are indented further than the section's baseline.
      3. Switching numbering definition (num_id) entirely for sub-items --
         a Word quirk where indenting a bullet without a properly configured
         second list level creates a whole new list instead of using ilvl=1
         of the same list. The first-seen num_id per section is treated as
         the "main" list; any different num_id is a sub-bullet.
    """
    if not items:
        return []

    if any(ilvl for _, ilvl, _, _ in items):
        return [{"text": t, "level": ilvl or 0} for t, ilvl, _, _ in items]

    indents = [li for _, _, li, _ in items if li is not None]
    if indents and len(set(indents)) > 1:
        baseline = min(indents)
        return [{"text": t, "level": 1 if (li or 0) > baseline else 0} for t, _, li, _ in items]

    num_ids = [nid for _, _, _, nid in items if nid is not None]
    if num_ids and len(set(num_ids)) > 1:
        # A numId switch only means "entering a sub-bullet block" when the
        # PRECEDING line ends with a colon (a natural "here are the details"
        # cue used throughout these documents). Without that, a numId switch
        # is just Word creating a new list instance for what is still a
        # top-level bullet -- e.g. a document can have a genuine main-bullet
        # numId, a genuine sub-bullet numId (following a colon), AND a THIRD
        # numId for one more independent main bullet at the very end. Simply
        # treating "any numId different from the first one" as sub-level
        # breaks that case, since it can't distinguish the second main
        # bullet from a real sub-item.
        result = []
        current_main_num_id = None
        sub_num_id = None
        prev_ends_with_colon = False
        for text, _, _, nid in items:
            if nid is None:
                level = 0
                sub_num_id = None
            elif current_main_num_id is None:
                current_main_num_id = nid
                level = 0
            elif nid == current_main_num_id:
                level = 0
                sub_num_id = None
            elif nid == sub_num_id:
                level = 1  # continuing an already-established sub-block
            elif prev_ends_with_colon:
                sub_num_id = nid
                level = 1  # a new sub-block, cued by the preceding colon
            else:
                current_main_num_id = nid  # a new main bullet, not a sub-item
                sub_num_id = None
                level = 0
            result.append({"text": text, "level": level})
            prev_ends_with_colon = text.rstrip().endswith(":")
        return result

    return [{"text": t, "level": 0} for t, _, _, _ in items]


def parse_arbeitspapier(docx_path_or_file):
    """
    Parses one Arbeitspapier working paper and returns a dict with the fields
    needed to populate a Control + ArbeitspapierSubmission:
        control_id, kontrollziel, kontrollbeschreibung,
        test_activities (JSON-encoded nested bullet structure),
        kontrollergebnis_raw, no_deviation, geprueft_von, geprueft_datum
    """
    doc = Document(docx_path_or_file)
    sections = _extract_sections(doc)
    geprueft_von, geprueft_datum = _extract_geprueft_metadata(doc)

    def plain_text(section_key):
        return " ".join(t for t, _, _, _ in sections.get(section_key, [])).strip()

    control_id_line = plain_text("1.1")
    # e.g. "S-CC-1.0 Business Conduct Guidelines" -> take the first token as the ID
    control_id = control_id_line.split(" ")[0] if control_id_line else ""

    kontrollziel = plain_text("1.2")
    kontrollbeschreibung = plain_text("1.3")

    # Section 5.1: resolve nesting robustly across the different ways working
    # papers encode sub-bullets (see _resolve_bullet_levels)
    activities = _resolve_bullet_levels(sections.get("5.1", []))
    test_activities = json.dumps(activities)

    kontrollergebnis_raw = plain_text("5.3")
    no_deviation = any(marker in kontrollergebnis_raw.lower() for marker in NO_DEVIATION_MARKERS)

    return {
        "control_id": control_id,
        "kontrollziel": kontrollziel,
        "kontrollbeschreibung": kontrollbeschreibung,
        "test_activities": test_activities,
        "kontrollergebnis_raw": kontrollergebnis_raw,
        "no_deviation": no_deviation,
        "geprueft_von": geprueft_von,
        "geprueft_datum": geprueft_datum,
    }


def compute_diff(old_values, new_values, fields):
    """
    Compares old_values and new_values dicts across the given field names
    and returns {"field": {"old": ..., "new": ...}} only for fields that
    actually changed. Used to populate ChangeLogEntry.diff.
    """
    diff = {}
    for field in fields:
        old = old_values.get(field, "")
        new = new_values.get(field, "")
        if old != new:
            diff[field] = {"old": old, "new": new}
    return diff


def import_arbeitspapier(project, file_obj, filename, user):
    """
    Parses one uploaded Arbeitspapier file and stores it as a Control +
    ArbeitspapierSubmission for the given project. If the control doesn't
    exist yet (e.g. Excel wasn't imported first), it's created from the
    working paper's own 1.2/1.3 sections.

    Runs the validation checks (see VALIDATORS above) comparing the working
    paper's Kontrollziel/Kontrollbeschreibung against whatever the master
    IKS Excel already established for this control, BEFORE this import
    potentially fills in/overwrites those fields.

    Returns the created ArbeitspapierSubmission.
    """
    parsed = parse_arbeitspapier(file_obj)

    if not parsed["control_id"]:
        raise ValueError(f"Could not find a control ID (section 1.1) in {filename}")

    control, _ = Control.objects.get_or_create(
        project=project,
        control_id=parsed["control_id"],
        defaults={
            "kontrollziel": parsed["kontrollziel"],
            "kontrollbeschreibung": parsed["kontrollbeschreibung"],
        },
    )

    try:
        # Snapshot the master values BEFORE we potentially fill them in below --
        # this is what the working paper gets validated against.
        control_before = {
            "kontrollziel": control.kontrollziel,
            "kontrollbeschreibung": control.kontrollbeschreibung,
        }
        validation_results = run_validations(control_before, parsed, project.language)

        # Fill in kontrollziel/beschreibung if missing (e.g. was blank from Excel import)
        changed = False
        if not control.kontrollziel and parsed["kontrollziel"]:
            control.kontrollziel = parsed["kontrollziel"]
            changed = True
        if not control.kontrollbeschreibung and parsed["kontrollbeschreibung"]:
            control.kontrollbeschreibung = parsed["kontrollbeschreibung"]
            changed = True
        if changed:
            control.save()

        # Diff against the previous latest submission (if any) for the change log --
        # this is what powers the "highlight what changed" view in the history page.
        previous = control.latest_submission
        previous_values = {
            "test_activities": previous.test_activities_as_editable_text() if previous else "",
            "kontrollergebnis_raw": previous.kontrollergebnis_raw if previous else "",
        }

        submission = ArbeitspapierSubmission.objects.create(
            control=control,
            uploaded_by=user,
            test_activities=parsed["test_activities"],
            kontrollergebnis_raw=parsed["kontrollergebnis_raw"],
            no_deviation=parsed["no_deviation"],
            validation_results=json.dumps(validation_results),
            geprueft_von=parsed["geprueft_von"],
            geprueft_datum=parsed["geprueft_datum"],
            source_file=filename,
        )

        new_values = {
            "test_activities": submission.test_activities_as_editable_text(),
            "kontrollergebnis_raw": submission.kontrollergebnis_raw,
        }
        diff = compute_diff(previous_values, new_values, ["test_activities", "kontrollergebnis_raw"])
        ChangeLogEntry.objects.create(
            control=control, event_type="import", actor=user, source_file=filename, diff=diff,
        )

        if control.import_error:
            control.import_error = ""
            control.save(update_fields=["import_error"])

        return submission
    except Exception as e:
        control.import_error = f"{filename}: {e}"
        control.save(update_fields=["import_error"])
        raise


# ---------------------------------------------------------------------------
# 3. Report generation (docxtpl)
# ---------------------------------------------------------------------------

def _group_activities(activities):
    """
    Groups the flat (text, level) list into a nested structure:
    [{"text": <main bullet text>, "sub_items": [<sub bullet text>, ...]}, ...]

    This lets the report template use two separate placeholder paragraphs
    (one for main bullets, one for sub-bullets) that you can format freely
    in Word -- indentation, spacing, bullet style, whatever you want --
    without touching any Python code.
    """
    grouped = []
    for item in activities:
        level = item.get("level", 0) or 0
        text = item.get("text", "")
        if level == 0:
            grouped.append({"text": text, "sub_items": []})
        else:
            if grouped:
                grouped[-1]["sub_items"].append(text)
            else:
                # Sub-bullet with no preceding main bullet (shouldn't normally
                # happen) -- treat it as its own main bullet instead of losing it.
                grouped.append({"text": text, "sub_items": []})
    return grouped


def resolve_template(project):
    """
    Finds the active ReportTemplate for this project's audit_type + language.
    Returns the ReportTemplate instance, or None if nothing's been uploaded yet.
    """
    from .models import ReportTemplate
    return ReportTemplate.objects.filter(
        audit_type=project.audit_type, language=project.language, is_active=True
    ).first()


def generate_report(project, template_path):
    """
    Renders the final Word report for a project using the given docxtpl
    template path. Returns an in-memory BytesIO of the .docx file.
    """
    controls = []
    for control in project.controls.all().order_by("control_id"):
        submission = control.latest_submission
        activities = submission.test_activities_list if submission else []
        controls.append({
            "control_id": control.control_id,
            "kontrollziel": control.kontrollziel,
            "kontrollbeschreibung": control.kontrollbeschreibung,
            "test_activities": _group_activities(activities),
            "result_text": submission.result_text if submission else "[Not yet assessed]",
        })

    audit_periods = [
        {
            "label": p.label or f"Period {i+1}",
            "start_date": p.start_date.strftime("%d %B %Y"),
            "end_date": p.end_date.strftime("%d %B %Y"),
        }
        for i, p in enumerate(project.audit_periods.all())
    ]

    context = {
        "customer_name": project.customer_name,
        "customer_address": project.customer_address,
        # Swapped per explicit request: {{ report_kind }} now outputs the
        # audit standard (SOC 2 / PS 951 / PS 3000), {{ audit_type }} now
        # outputs Type 1 / Type 2. The underlying model fields/form dropdowns
        # keep their original meaning -- only what these two template keys
        # render has changed.
        "report_kind": project.get_audit_type_display(),
        "audit_type": project.get_report_kind_display(),
        "report_date": format_custom_date(project.report_date, "de_ordinal" if project.language == "de" else "en_ordinal"),
        "is_type2": project.is_type2,
        "audit_periods": audit_periods,
        "controls": controls,
    }
    # Custom placeholders (defined on the Placeholders page, values filled in
    # per-project) are injected directly at the top level -- {{ key }} in
    # templates, not {{ extra.key }} -- so they work exactly like built-in
    # placeholders from the template author's point of view. Date-type ones
    # are auto-formatted per the project's report language (English ordinal
    # for "en", German for "de") from the raw ISO date stored in
    # Project.extra_fields -- no manual per-placeholder format choice needed.
    from .models import Placeholder
    import datetime
    extra = dict(project.extra_fields or {})
    date_keys = set(Placeholder.objects.filter(is_custom=True, field_type="date").values_list("key", flat=True))
    date_format = "de_ordinal" if project.language == "de" else "en_ordinal"
    for key, value in extra.items():
        if key in date_keys and value:
            try:
                date_obj = datetime.date.fromisoformat(value)
                extra[key] = format_custom_date(date_obj, date_format)
            except (ValueError, TypeError):
                pass  # leave raw value as-is if it's not a valid ISO date
    context.update(extra)

    # audit_conducted_from_to (Type 2 only) is computed from two raw dates
    # stored in extra_fields, not a plain value itself -- e.g.
    # "1st January 2026 to 30th June 2026" (EN) / "1. Januar 2026 bis 30. Juni 2026" (DE)
    from_raw = (project.extra_fields or {}).get("audit_conducted_from", "")
    to_raw = (project.extra_fields or {}).get("audit_conducted_to", "")
    audit_conducted_from_to = ""
    if from_raw and to_raw:
        try:
            from_date = datetime.date.fromisoformat(from_raw)
            to_date = datetime.date.fromisoformat(to_raw)
            joiner = "bis" if project.language == "de" else "to"
            audit_conducted_from_to = f"{format_custom_date(from_date, date_format)} {joiner} {format_custom_date(to_date, date_format)}"
        except (ValueError, TypeError):
            pass
    context["audit_conducted_from_to"] = audit_conducted_from_to

    tpl = DocxTemplate(template_path)
    tpl.render(context)

    buffer = io.BytesIO()
    tpl.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# 4. Report template management (upload, preview conversion, placeholder scan)
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(input_path, output_dir):
    """
    Converts a .docx to .pdf for in-browser template preview.

    Tries headless LibreOffice first (best fidelity -- preserves table cell
    colors, exact fonts, complex formatting). If LibreOffice isn't available
    on this machine, falls back to a pure-Python route (mammoth + weasyprint)
    -- no LibreOffice/MS Word needed, but LOWER FIDELITY: table cell
    background colors and some complex formatting are lost in that path,
    though text, structure, and basic table layout come through fine.

    Returns the path to the generated PDF, or None if both routes failed.
    """
    result = _convert_via_libreoffice(input_path, output_dir)
    if result:
        return result
    return _convert_via_mammoth_weasyprint(input_path, output_dir)


def _convert_via_libreoffice(input_path, output_dir):
    import subprocess
    import shutil
    import tempfile
    import uuid

    candidates = [
        "soffice",
        "libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS (Homebrew cask / direct download)
        "/usr/bin/soffice",       # common Linux path (incl. PythonAnywhere, pre-installed since ~2020)
        "/usr/local/bin/soffice",  # common Linux/Homebrew-on-Intel-Mac path
        "/opt/homebrew/bin/soffice",  # Homebrew on Apple Silicon, if symlinked
    ]

    soffice_bin = None
    for candidate in candidates:
        if "/" in candidate:
            if Path(candidate).exists():
                soffice_bin = candidate
                break
        elif shutil.which(candidate):
            soffice_bin = candidate
            break

    if not soffice_bin:
        return None

    # Each invocation gets its own LibreOffice user profile directory.
    # Without this, concurrent requests (e.g. two people clicking "Preview"
    # around the same time on a shared host) share the same default profile
    # and can crash or hang each other -- LibreOffice's headless mode isn't
    # safe for concurrent access to a single profile.
    user_install_dir = Path(tempfile.gettempdir()) / f"libreoffice_profile_{uuid.uuid4().hex}"
    try:
        subprocess.run(
            [
                soffice_bin, "--headless", "--norestore",
                f"-env:UserInstallation=file://{user_install_dir}",
                "--convert-to", "pdf", "--outdir", str(output_dir), str(input_path),
            ],
            check=True, capture_output=True, timeout=60,
        )
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return None
    finally:
        # Each profile directory is used exactly once and can be sizeable --
        # clean it up so repeated conversions don't slowly fill up disk quota.
        shutil.rmtree(user_install_dir, ignore_errors=True)

    pdf_path = Path(output_dir) / (Path(input_path).stem + ".pdf")
    return pdf_path if pdf_path.exists() else None


def _convert_via_mammoth_weasyprint(input_path, output_dir):
    """
    LibreOffice-free fallback. Requires `pip install mammoth weasyprint`.
    weasyprint itself needs Pango/cairo system libraries (much lighter than
    LibreOffice: `brew install pango` on macOS, or
    `apt install libpango-1.0-0 libpangocairo-1.0-0 libgdk-pixbuf2.0-0` on
    Linux) -- not literally zero system dependencies, but far smaller.
    """
    try:
        import mammoth
        from weasyprint import HTML
    except ImportError:
        return None

    try:
        with open(input_path, "rb") as docx_file:
            html = mammoth.convert_to_html(docx_file).value
        pdf_path = Path(output_dir) / (Path(input_path).stem + ".pdf")
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        HTML(string=f"<html><body>{html}</body></html>").write_pdf(str(pdf_path))
        return pdf_path if pdf_path.exists() else None
    except Exception:
        return None


PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w\.]+)\s*\}\}")
FOR_LOOP_RE = re.compile(r"\{%\s*(?:p|tr|tc|r)?\s*for\s+(\w+)\s+in\s+([\w\.]+)")
IF_RE = re.compile(r"\{%\s*if\s+(?:not\s+)?([\w\.]+)")


def _all_text_blocks(doc):
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def resolve_documented_ancestor(key, loop_bindings, catalog_keys):
    """
    Given a placeholder key that may be a loop variable (e.g. "s" from
    "{%p for s in m.sub_items %}", where "m" is itself a loop variable from
    "{%p for m in c.test_activities %}"), walks up the loop chain looking
    for the nearest ancestor path that's actually a documented catalog key.

    Checks prefix matches at EVERY step before substituting further, since
    the catalog documents things using the loop variable's own name (e.g.
    "c.test_activities") rather than the fully-resolved root (e.g.
    "controls.test_activities") -- substituting all the way to the root
    would miss the actual documented entry.

    Returns the matching catalog key, or None if nothing in the chain matches.
    """
    current = key
    seen = set()
    for _ in range(6):  # generous depth cap against malformed/cyclic templates
        parts = current.split(".")
        for i in range(len(parts), 0, -1):
            candidate = ".".join(parts[:i])
            if candidate in catalog_keys:
                return candidate
        root = parts[0]
        if root in loop_bindings and root not in seen:
            seen.add(root)
            rest = current[len(root):]
            current = loop_bindings[root] + rest
        else:
            break
    return None


def scan_template_placeholders(docx_path):
    """
    Scans every paragraph and table cell in a .docx for {{ placeholder }}
    tags, {% for x in y %} loop bindings, and {% if x %} conditionals.

    Returns {"placeholders": [...], "loop_bindings": {...}, "if_conditions": {...}}
    Each placeholder dict: {key, root, is_loop_var, loop_source}
    - key: the raw placeholder as found, e.g. "m.text"
    - root: the leading variable name, e.g. "m"
    - is_loop_var: True if `root` is bound by an enclosing for-loop
    - loop_source: what it loops over, e.g. "c.test_activities" (if is_loop_var)

    Catalog-aware resolution (which ancestor a loop variable documents to)
    is deliberately NOT done here -- see resolve_documented_ancestor(), called
    by the view with the actual catalog, keeping this function purely mechanical.
    """
    doc = Document(docx_path)
    text_blocks = list(_all_text_blocks(doc))

    # First pass: collect all loop variable bindings (var_name -> source expression)
    loop_bindings = {}
    if_conditions = set()
    for text in text_blocks:
        for var_name, source_expr in FOR_LOOP_RE.findall(text):
            loop_bindings[var_name] = source_expr
        for cond_var in IF_RE.findall(text):
            if_conditions.add(cond_var)

    # Second pass: collect all {{ }} placeholder usages
    found_keys = set()
    for text in text_blocks:
        for key in PLACEHOLDER_RE.findall(text):
            found_keys.add(key)

    results = []
    for key in sorted(found_keys):
        root = key.split(".")[0]
        is_loop_var = root in loop_bindings
        results.append({
            "key": key, "root": root, "is_loop_var": is_loop_var,
            "loop_source": loop_bindings.get(root) if is_loop_var else None,
        })

    return {"placeholders": results, "loop_bindings": loop_bindings, "if_conditions": if_conditions}


# ---------------------------------------------------------------------------
# 5. Word-level diff rendering for the change history page
# ---------------------------------------------------------------------------

def compute_word_diff_html(old_text, new_text):
    """
    Produces one inline HTML fragment showing a word-level diff between
    old_text and new_text -- deleted words struck through in red, added
    words highlighted in green, unchanged text in between left plain.

    Uses Google's diff-match-patch. Text is split on word boundaries first
    (rather than diffing raw characters) so the result reads as whole-word
    changes ("policyy" -> "policy" shows just the extra "y" removed, not
    every character re-flagged), then diff_cleanupSemantic groups the
    result into natural phrase-sized chunks.
    """
    from diff_match_patch import diff_match_patch
    import html as html_module
    import re as re_module

    dmp = diff_match_patch()

    # diff-match-patch's line-mode diffing works well for our multi-line bullet
    # text: chars_to_lines maps each whole line to a single token so the diff
    # operates at line granularity, THEN we refine within changed lines at the
    # word level for a tighter, more readable result.
    a_words = re_module.findall(r"\S+|\s+", old_text or "")
    b_words = re_module.findall(r"\S+|\s+", new_text or "")

    # diff_match_patch works on strings; encode each unique "word" (or run of
    # whitespace) as a private-use-area character so word-level tokens survive
    # the diff as atomic units instead of being diffed character-by-character.
    word_to_char = {}
    char_to_word = []

    def encode(words):
        chars = []
        for w in words:
            if w not in word_to_char:
                word_to_char[w] = chr(0xE000 + len(char_to_word))
                char_to_word.append(w)
            chars.append(word_to_char[w])
        return "".join(chars)

    a_encoded = encode(a_words)
    b_encoded = encode(b_words)

    diffs = dmp.diff_main(a_encoded, b_encoded, checklines=False)
    dmp.diff_cleanupSemantic(diffs)

    parts = []
    for op, encoded_chunk in diffs:
        text = "".join(char_to_word[ord(c) - 0xE000] for c in encoded_chunk)
        escaped = html_module.escape(text)
        if op == 0:
            parts.append(escaped)
        elif op == -1:
            parts.append(f'<span style="background:#fcedee; color:#b32d2e; text-decoration:line-through;">{escaped}</span>')
        elif op == 1:
            parts.append(f'<span style="background:#edfaef; color:#00a32a;">{escaped}</span>')

    return "".join(parts).replace("\n", "<br>")


# ---------------------------------------------------------------------------
# 6. Reference report import (old finished reports, for reviewer reference)
# ---------------------------------------------------------------------------

# Header text patterns identifying the results table, matched case-insensitively
# as substrings -- flexible on exact wording/language since old reports vary.
RESULT_HEADER_PATTERNS = ["testergebnis der kontrolle", "result of test"]
TEST_PERFORMED_HEADER_PATTERNS = ["performed by rsm ebner", "durch rsm ebner"]

# A control ID looks like "S-CC-1.0", "IKS-014", etc: short, no spaces, has a
# letter+digit or hyphen pattern. Used to split "S-CC-1.0\nDescription..." into
# (control_id, description) when the first line of the description cell looks
# like an ID rather than prose.
CONTROL_ID_LINE_RE = re.compile(r"^[A-Za-z0-9]+[-_][A-Za-z0-9.\-]+$")


def _extract_cell_bullets_as_dash_text(cell):
    """
    Reads a table cell's paragraphs and returns them as plain text with the
    same "- " sub-bullet convention used for Arbeitspapier test activities
    elsewhere in the app: main bullets as-is, sub-bullets prefixed with "- ".

    Sub-bullets are detected two ways, since different documents encode them
    differently:
      1. Proper Word multilevel lists (numPr.ilvl > 0) -- used by e.g. the
         Arbeitspapier working papers and some old reports.
      2. Plain paragraph indentation only, with ilvl staying 0 throughout --
         seen in some reports, where sub-items are just indented further
         (a larger pPr.ind.left) than the main bullets in the same cell,
         with no real multilevel list structure at all.

    Falls back to the cell's raw text if it has no paragraphs at all.
    """
    paras_info = []
    for p in cell.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        ilvl = None
        left_indent = 0
        pPr = p._p.pPr
        if pPr is not None:
            if pPr.numPr is not None and pPr.numPr.ilvl is not None:
                ilvl = pPr.numPr.ilvl.val
            if pPr.ind is not None and pPr.ind.left is not None:
                left_indent = pPr.ind.left
        paras_info.append({"text": text, "ilvl": ilvl, "left_indent": left_indent})

    if not paras_info:
        return cell.text.strip()

    baseline_indent = min(p["left_indent"] for p in paras_info)

    lines = []
    for p in paras_info:
        is_sub = bool(p["ilvl"]) or p["left_indent"] > baseline_indent
        prefix = "- " if is_sub else ""
        lines.append(f"{prefix}{p['text']}")
    return "\n".join(lines)


def _find_results_tables(doc):
    """
    Scans every table in the document for a header row matching the results
    table pattern (result column on the far right, test-performed column
    just to its left). Returns a list of (table, header_row_index,
    kontrollbeschreibung_col_index) for every match found.
    """
    matches = []
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 3:
                continue
            last_text = cells[-1].text.strip().lower()
            second_last_text = cells[-2].text.strip().lower()
            is_result_header = any(p in last_text for p in RESULT_HEADER_PATTERNS)
            is_test_header = any(p in second_last_text for p in TEST_PERFORMED_HEADER_PATTERNS)
            if is_result_header and is_test_header:
                kontroll_col_index = len(cells) - 3  # column just left of "tests performed"
                matches.append((table, row_idx, kontroll_col_index))
                break  # one header row per table is enough
    return matches


def parse_reference_report(docx_path_or_file):
    """
    Parses an old/finished report's results table(s) into a list of dicts:
    [{"criteria", "control_id", "kontrollbeschreibung", "test_performed", "result_text"}, ...]
    Returns an empty list if no matching table is found.
    """
    doc = Document(docx_path_or_file)
    table_matches = _find_results_tables(doc)

    entries = []
    for table, header_row_idx, kontroll_col in table_matches:
        for row in table.rows[header_row_idx + 1:]:
            cells = row.cells
            if len(cells) < 3:
                continue

            kontroll_cell_text = cells[kontroll_col].text.strip() if kontroll_col >= 0 else ""
            test_performed = _extract_cell_bullets_as_dash_text(cells[-2])
            result_text = cells[-1].text.strip()
            criteria = cells[0].text.strip() if kontroll_col > 0 else ""

            if not kontroll_cell_text and not test_performed and not result_text:
                continue  # skip fully blank rows

            # Split "S-CC-1.0\nDescription..." into control_id + description
            # if the first line looks like an ID, otherwise keep it all as description.
            lines = kontroll_cell_text.split("\n", 1)
            if len(lines) == 2 and CONTROL_ID_LINE_RE.match(lines[0].strip()):
                control_id, kontrollbeschreibung = lines[0].strip(), lines[1].strip()
            else:
                control_id, kontrollbeschreibung = "", kontroll_cell_text

            entries.append({
                "criteria": criteria,
                "control_id": control_id,
                "kontrollbeschreibung": kontrollbeschreibung,
                "test_performed": test_performed,
                "result_text": result_text,
            })

    return entries


def import_reference_report(file_obj, filename, user):
    """
    Parses an uploaded old report and stores it + its entries in the DB.
    Returns the created ReferenceReport. Raises ValueError if no matching
    results table was found in the document.
    """
    from .models import ReferenceReport, ReferenceReportEntry

    entries = parse_reference_report(file_obj)
    if not entries:
        raise ValueError(
            "Couldn't find a results table in this document (looking for a table with a "
            '"Result of test" / "Testergebnis der Kontrolle" column). '
            "Nothing was imported."
        )

    file_obj.seek(0)  # parsing above consumed the stream -- rewind before saving the actual file
    report = ReferenceReport.objects.create(name=filename, file=file_obj, uploaded_by=user)
    for i, entry in enumerate(entries):
        ReferenceReportEntry.objects.create(reference_report=report, order=i, **entry)

    rebuild_reference_fts_index()
    return report


# ---------------------------------------------------------------------------
# 7. "Similar past entries" suggestions for section 5.1, sourced from the
#    Reference Report library (SQLite FTS5 + trigram tokenizer + BM25 rank,
#    with a DB-agnostic fallback for Postgres in production)
# ---------------------------------------------------------------------------

STOPWORDS_EN = {
    "the", "a", "an", "of", "for", "and", "or", "to", "in", "on", "with", "by",
    "is", "are", "was", "were", "this", "that", "these", "those", "as", "at",
    "be", "from", "it", "its", "which", "we", "our", "their", "been", "not",
    "no", "all", "any", "such", "into", "will", "shall", "has", "have", "had",
}
STOPWORDS_DE = {
    "der", "die", "das", "und", "oder", "zu", "zur", "zum", "für", "mit",
    "von", "bei", "im", "in", "auf", "des", "eine", "einer", "eines", "ist",
    "sind", "wurde", "wurden", "als", "auch", "nicht", "werden", "wird",
    "sich", "dass", "durch", "den", "dem", "ein", "einem", "einen", "sowie",
}
STOPWORDS = STOPWORDS_EN | STOPWORDS_DE

WORD_RE = re.compile(r"[a-zA-ZäöüÄÖÜß]+")


# ---------------------------------------------------------------------------
# Bilingual (DE<->EN) glossary for cross-language suggestion matching.
#
# Pure keyword-overlap scoring can't match a German query against an English
# reference entry (or vice versa) -- "Schulung" and "training" share no
# letters. Rather than reaching for a heavy embeddings model (external
# dependency, cost, breaks the "runs on a $4 droplet" design goal), this is a
# curated glossary of common IKS/SOC2/ISAE audit vocabulary: when a query
# keyword has a known translation, the translation is added to the keyword
# set too, so the SAME word-overlap scoring engine can credit a match across
# languages. Not exhaustive -- covers the vocabulary that actually recurs in
# audit working papers/reports. Extend freely as new gaps show up.
# ---------------------------------------------------------------------------
GLOSSARY_DE_EN = {
    # Training / awareness
    "schulung": "training", "schulungen": "training", "schulungsprogramm": "training program",
    "schulungsinhalte": "training content", "sensibilisierung": "awareness", "sensibilisierungs": "awareness",
    "training": "schulung", "awareness": "sensibilisierung",
    # Security / policy / governance
    "sicherheit": "security", "security": "sicherheit",
    "richtlinie": "policy", "richtlinien": "policy", "leitlinie": "guideline", "leitlinien": "guideline",
    "policy": "richtlinie", "guideline": "leitlinie",
    "rahmenwerk": "framework", "framework": "rahmenwerk",
    "governance": "steuerung", "steuerung": "governance control",
    "vorgabe": "requirement standard", "vorgaben": "requirements",
    "standard": "standard",
    # Process / people
    "prozessverantwortlichen": "process owner", "prozessverantwortlicher": "process owner",
    "verantwortlichen": "responsible owner", "owner": "verantwortlicher",
    "mitarbeiter": "employee", "mitarbeitern": "employee", "employee": "mitarbeiter", "employees": "mitarbeiter",
    "fachabteilung": "department business unit", "department": "fachabteilung",
    "geschaeftsleitung": "management executive", "geschäftsleitung": "management", "management": "geschäftsleitung",
    # Review / audit / test
    "pruefung": "review test audit examination", "prüfung": "review test audit examination",
    "review": "prüfung", "audit": "prüfung revision", "examination": "prüfung", "test": "prüfung test",
    "revision": "audit internal audit", "interne revision": "internal audit",
    "einsichtnahme": "inspection review", "inspection": "einsichtnahme",
    "durchfuehrung": "execution performance", "durchführung": "execution", "execution": "durchführung",
    "stichprobe": "sample", "sample": "stichprobe",
    "pruefungszeitraum": "audit period", "prüfungszeitraum": "audit period",
    "feststellung": "finding", "feststellungen": "finding", "finding": "feststellung",
    "abweichung": "deviation", "abweichungen": "deviation", "deviation": "abweichung", "deviations": "abweichung",
    "gespraech": "discussion", "gespräch": "discussion", "discussion": "gespräch",
    "nachvollzug": "review verification", "verification": "nachvollzug",
    # Documentation
    "dokumentation": "documentation", "documentation": "dokumentation",
    "dokument": "document", "document": "dokument",
    "prozessdokumentation": "process documentation", "prozessbeschreibung": "process description",
    "handbuch": "manual handbook", "manual": "handbuch", "handbook": "handbuch",
    # Risk
    "risiko": "risk", "risk": "risiko", "risiken": "risks",
    "risikomanagement": "risk management", "risikoanalyse": "risk assessment analysis",
    "risk assessment": "risikoanalyse", "risk management": "risikomanagement",
    # Access / authorization
    "zugriff": "access", "access": "zugriff", "zutritt": "physical access", "zutritts": "access",
    "zutritten": "access", "berechtigung": "authorization permission access",
    "berechtigungen": "authorization", "authorization": "berechtigung", "permission": "berechtigung",
    "permissions": "berechtigung", "berechtigungsmanagement": "access management",
    "access management": "berechtigungsmanagement", "rollenmatrix": "role matrix",
    "rollenkonzept": "role concept rbac", "role": "rolle", "rolle": "role",
    "funktionstrennung": "segregation of duties", "segregation of duties": "funktionstrennung",
    "vier-augen-prinzip": "four eyes principle dual control", "vier augen prinzip": "four eyes principle",
    "kennwort": "password", "passwort": "password", "password": "kennwort passwort",
    "authentifizierung": "authentication", "authentication": "authentifizierung",
    "multi-faktor": "multi factor mfa", "mfa": "multi-faktor",
    # Vulnerability / security operations
    "schwachstelle": "vulnerability", "schwachstellen": "vulnerability vulnerabilities",
    "vulnerability": "schwachstelle", "vulnerabilities": "schwachstellen",
    "sicherheitsaudit": "security audit", "sicherheitsaudits": "security audits",
    "sicherheitsvorfall": "security incident", "security incident": "sicherheitsvorfall",
    "vorfall": "incident", "incident": "vorfall", "vorfaelle": "incidents", "vorfälle": "incidents",
    "hardening": "systemhärtung", "systemhaertung": "hardening", "systemhärtung": "hardening", "haertung": "hardening",
    "patch": "patch", "patches": "patches", "patch management": "patch-management",
    "endpoint": "endgerät", "endgeraet": "endpoint", "endgerät": "endpoint",
    "malware": "schadcode", "schadcode": "malware", "virenschutz": "antivirus malware protection",
    # Change / configuration
    "aenderung": "change", "änderung": "change", "change": "änderung", "aenderungen": "changes",
    "änderungen": "changes", "aenderungsmanagement": "change management", "änderungsmanagement": "change management",
    "change management": "änderungsmanagement", "konfiguration": "configuration", "configuration": "konfiguration",
    "genehmigung": "approval", "freigabe": "approval release", "approval": "freigabe genehmigung",
    # Backup / continuity
    "datensicherung": "backup data protection", "sicherung": "backup", "backup": "datensicherung sicherung",
    "wiederherstellung": "recovery restore", "recovery": "wiederherstellung", "restore": "wiederherstellung",
    "archivierung": "archiving", "archiving": "archivierung",
    "notfall": "emergency disaster", "business continuity": "geschäftsfortführung",
    "geschaeftsfortfuehrung": "business continuity", "geschäftsfortführung": "business continuity",
    "notfallmanagement": "disaster recovery emergency management",
    # Monitoring / logging
    "ueberwachung": "monitoring", "überwachung": "monitoring", "monitoring": "überwachung",
    "protokollierung": "logging", "protokoll": "log logging", "logging": "protokollierung",
    "kapazitaet": "capacity", "kapazität": "capacity", "capacity": "kapazität",
    "kapazitaeten": "capacities", "kapazitäten": "capacities",
    # Physical security
    "zutrittskontrolle": "physical access control", "physische sicherheit": "physical security",
    "physical security": "physische sicherheit", "rechenzentrum": "data center", "data center": "rechenzentrum",
    "standort": "site location", "standorte": "sites", "site": "standort",
    # Supplier / contracts
    "lieferant": "supplier vendor", "lieferanten": "suppliers", "dienstleister": "service provider vendor",
    "supplier": "lieferant dienstleister", "vendor": "lieferant", "geschaeftspartnerkodex": "business partner code of conduct",
    "vertrag": "contract", "vertraege": "contracts", "verträge": "contracts", "contract": "vertrag",
    # Encryption / cryptography
    "verschluesselung": "encryption", "verschlüsselung": "encryption", "encryption": "verschlüsselung",
    "kryptographie": "cryptography", "cryptography": "kryptographie",
    # Compliance / certification
    "zertifizierung": "certification", "certification": "zertifizierung", "zertifikat": "certificate",
    "certificate": "zertifikat", "compliance": "compliance konformität", "konformitaet": "compliance",
    "konformität": "compliance",
    # Confidentiality/integrity/availability triad
    "vertraulichkeit": "confidentiality", "confidentiality": "vertraulichkeit",
    "integritaet": "integrity", "integrität": "integrity", "integrity": "integrität",
    "verfuegbarkeit": "availability", "verfügbarkeit": "availability", "availability": "verfügbarkeit",
    # Project / planning / strategy
    "projekt": "project", "project": "projekt", "planung": "planning", "planning": "planung",
    "strategie": "strategy", "strategy": "strategie", "ziel": "objective goal", "ziele": "objectives goals",
    "objective": "ziel", "goal": "ziel", "kommunikation": "communication", "communication": "kommunikation",
    # Ticket / process infrastructure
    "ticket": "ticket", "tickets": "tickets", "itsm": "itsm",
    "kunde": "customer client", "kunden": "customers", "customer": "kunde",
    "onboarding": "onboarding", "abwesenheit": "absence", "absence": "abwesenheit",
    "netzwerk": "network", "network": "netzwerk", "firewall": "firewall",
    "jaehrlich": "annual annually", "jährlich": "annual annually", "annual": "jährlich", "annually": "jährlich",
    "quartalsweise": "quarterly", "quarterly": "quartalsweise", "halbjaehrlich": "semi-annual", "halbjährlich": "semi-annual",
    "verpflichtend": "mandatory", "mandatory": "verpflichtend", "kontinuierlich": "continuous", "continuous": "kontinuierlich",
    "wartung": "maintenance", "maintenance": "wartung",
    "ethik": "ethics", "ethics": "ethik", "integritaet und ethik": "integrity and ethics",
    "asset": "asset", "assets": "assets", "lebenszyklus": "lifecycle", "lifecycle": "lebenszyklus",
}


def _translate_keywords(keywords):
    """Expands a keyword list with any known glossary translations (DE<->EN)."""
    expanded = list(keywords)
    for kw in keywords:
        translation = GLOSSARY_DE_EN.get(kw)
        if translation:
            expanded.extend(translation.split())
    return expanded


def extract_keywords(text, top_n=12):
    """
    Splits text into words, drops stopwords and very short words, and returns
    the top_n most frequent remaining words -- the "meaningful terms" used to
    drive the similarity search, per the stopword-filtering + frequency
    approach. Also expands long words with any embedded compound roots (see
    _expand_with_compound_roots) so German compounds contribute their
    meaningful root as a matchable term too.
    """
    from collections import Counter
    words = WORD_RE.findall((text or "").lower())
    filtered = [w for w in words if w not in STOPWORDS and len(w) > 2]
    filtered = _expand_with_compound_roots(filtered)
    counts = Counter(filtered)
    return [w for w, _ in counts.most_common(top_n)]


def rebuild_reference_fts_index():
    """
    Drops and rebuilds the SQLite FTS5 (trigram tokenizer) index over all
    ReferenceReportEntry text. No-op on non-SQLite backends -- suggestions
    fall back to a pure-Python keyword-overlap search there instead (see
    suggest_similar_test_activities()).

    Rebuilt from scratch on every reference-report import/delete rather than
    incrementally synced: the dataset (a handful of imported reports) is
    small enough that a full rebuild is cheap, and it avoids any risk of the
    index silently drifting out of sync with the real table.
    """
    from django.db import connection
    from .models import ReferenceReportEntry

    if connection.vendor != "sqlite":
        return

    with connection.cursor() as cursor:
        cursor.execute("DROP TABLE IF EXISTS reference_entry_fts")
        cursor.execute(
            "CREATE VIRTUAL TABLE reference_entry_fts USING fts5("
            "content, entry_id UNINDEXED, tokenize='trigram')"
        )
        rows = ReferenceReportEntry.objects.all().values_list("id", "kontrollbeschreibung", "test_performed")
        for entry_id, kb, tp in rows:
            cursor.execute(
                "INSERT INTO reference_entry_fts (content, entry_id) VALUES (%s, %s)",
                [f"{kb}\n{tp}", entry_id],
            )


def _entry_to_suggestion(entry, score):
    return {
        "entry_id": entry.id,
        "report_name": entry.reference_report.name,
        "control_id": entry.control_id,
        "kontrollbeschreibung": entry.kontrollbeschreibung,
        "test_performed": entry.test_performed,
        "score": round(float(score), 3) if score is not None else None,
    }


def _keywords_to_trigram_query(keywords, max_trigrams=60):
    """
    Decomposes each keyword into its own overlapping 3-character trigrams and
    OR's them all together. This is the actual point of trigram tokenization:
    quoting whole keywords as phrases (the previous approach) forces a
    near-exact substring match, which defeats fuzzy matching entirely --
    e.g. "kapazitätsmanagement" wouldn't match "Kapazitätsmonitoring" at all
    despite sharing a long common prefix. OR-ing individual trigrams lets
    BM25 reward entries that share MANY trigrams with the query (genuine
    lexical/topical overlap) even when word suffixes differ.
    """
    trigrams, seen = [], set()
    for kw in keywords:
        for i in range(len(kw) - 2):
            tri = kw[i:i + 3]
            if tri not in seen:
                seen.add(tri)
                trigrams.append(tri)
    trigrams = trigrams[:max_trigrams]
    return " OR ".join(f'"{t}"' for t in trigrams)


def _suggest_via_fts5(keywords, top_n):
    from django.db import connection
    from .models import ReferenceReportEntry

    query = _keywords_to_trigram_query(keywords)
    with connection.cursor() as cursor:
        try:
            cursor.execute(
                "SELECT entry_id, bm25(reference_entry_fts) as score "
                "FROM reference_entry_fts WHERE reference_entry_fts MATCH %s "
                "ORDER BY score LIMIT %s",
                [query, top_n * 5],  # over-fetch a bit; de-duped/trimmed below
            )
            rows = cursor.fetchall()
        except Exception:
            # e.g. FTS5/trigram not available in this SQLite build -- degrade gracefully
            return _suggest_via_keyword_overlap(keywords, top_n)

    entries_by_id = {
        e.id: e for e in ReferenceReportEntry.objects.filter(id__in=[r[0] for r in rows]).select_related("reference_report")
    }
    results, seen = [], set()
    for entry_id, score in rows:
        if entry_id in seen or entry_id not in entries_by_id:
            continue
        seen.add(entry_id)
        results.append(_entry_to_suggestion(entries_by_id[entry_id], score))
        if len(results) >= top_n:
            break
    return results


def _suggest_via_keyword_overlap(keywords, top_n):
    """DB-agnostic fallback (used on Postgres) -- scores by keyword overlap count."""
    from .models import ReferenceReportEntry

    keyword_set = set(keywords)
    scored = []
    for entry in ReferenceReportEntry.objects.all().select_related("reference_report"):
        combined = f"{entry.kontrollbeschreibung} {entry.test_performed}".lower()
        entry_words = set(WORD_RE.findall(combined))
        overlap = len(keyword_set & entry_words)
        if overlap:
            scored.append((overlap, entry))
    scored.sort(key=lambda x: -x[0])
    return [_entry_to_suggestion(entry, -overlap) for overlap, entry in scored[:top_n]]


def _shared_prefix_len(a, b):
    """Length of the common leading substring of two words."""
    n = 0
    for x, y in zip(a, b):
        if x != y:
            break
        n += 1
    return n


# Curated list of important compound roots for German. Words are frequently
# built by prepending modifiers to a head noun (e.g. "Informationssicherheits-
# risikomanagement" = "information security" + "risikomanagement"), so the
# meaningful root can be buried mid-word rather than at the start -- a plain
# prefix check misses it entirely. Rather than generic substring matching
# (tried, reverted: too noisy, degraded other cases), this only fires for a
# small curated list of terms that actually matter for audit/IKS content.
COMPOUND_ROOTS = [
    "risiko", "kapazität", "schwachstelle", "sicherheit", "audit", "kontrolle",
    "berechtigung", "authentifizierung", "verschlüsselung", "dokumentation",
    "prüfung", "überwachung", "schulung", "training", "wartung", "änderung",
    "backup", "sicherung", "wiederherstellung", "zugriff", "zertifizierung",
    "compliance", "management", "richtlinie", "prozess", "incident", "vorfall",
    "notfall", "kontinuität", "monitoring", "protokoll", "konfiguration",
    "revision", "lieferant", "vertrag", "asset", "patch", "netzwerk",
    "schad", "viren", "malware",
]

_UMLAUT_FOLD = str.maketrans({"ä": "a", "ö": "o", "ü": "u", "ß": "s"})


def _fold_umlauts(s):
    """
    Normalizes German umlauts/ß for substring comparison only (not used for
    display). Needed because German plurals often mutate the root vowel --
    e.g. "Vorfall" -> "Vorfälle"/"Vorfällen" -- so a literal substring check
    for "vorfall" inside "sicherheitsvorfällen" fails without this folding.
    """
    return s.translate(_UMLAUT_FOLD)


def _expand_with_compound_roots(words):
    """
    Given a list of words, returns them plus any curated COMPOUND_ROOTS found
    embedded within longer words (e.g. "informationssicherheitsrisikomanagement"
    also contributes "risiko" and "management" as extra matchable terms).
    Applied identically to both query keywords and reference-entry word
    counts, so a match on the shared root works as a normal exact match.
    Matching is done on umlaut-folded forms so plural mutations (e.g.
    "Vorfall" -> "Vorfällen") are still recognized as containing the root.
    """
    expanded = list(words)
    for word in words:
        if len(word) < 7:
            continue  # only worth checking words long enough to plausibly be a compound
        folded_word = _fold_umlauts(word)
        for root in COMPOUND_ROOTS:
            if root != word and _fold_umlauts(root) in folded_word:
                expanded.append(root)
    return expanded


def _best_partial_match_len(kw, w):
    """
    Length of the shared prefix between two words (handles compound-word
    suffix variants, e.g. "Kapazitätsmanagement" vs "Kapazitätsmonitoring" --
    same root, different ending). Requires both words to be >=6 chars.

    Note: substring-containment matching (catching e.g. "risikomanagement"
    embedded inside "Informationssicherheitsrisikomanagement") was tried
    here and reverted -- it introduced enough false-positive noise into the
    IDF calculation that it net-degraded other, previously-correct cases
    more than it fixed this one. Left as a known gap for now rather than
    risk destabilizing everything else.
    """
    if len(kw) < 6 or len(w) < 6:
        return 0
    return _shared_prefix_len(kw, w)


# Field weight multipliers: text found in/matching Test Activities (5.1)
# carries more signal than Kontrollbeschreibung, which often just describes
# generic background process context rather than what was actually done.
QUERY_FIELD_WEIGHT = {"test_activities": 2.0, "kontrollbeschreibung": 1.0}
ENTRY_FIELD_WEIGHT = {"test_performed": 2.0, "kontrollbeschreibung": 1.0}


def suggest_similar_test_activities(test_activities_text, kontrollbeschreibung_text="", top_n=5):
    """
    Given the current control's (in-progress) test activities text and its
    Kontrollbeschreibung, returns the top_n most similar past entries from
    the Reference Report library.

    Matches are weighted higher when they involve Test Activities (5.1),
    on BOTH sides:
      - Query side: keywords from YOUR control's test_activities count more
        than keywords from its Kontrollbeschreibung (which is often generic
        background process description, not what was actually tested).
      - Reference side: a match landing in a past entry's actual
        test_performed text is a stronger signal than one landing in that
        entry's Kontrollbeschreibung.
    A keyword appearing in both your test_activities AND Kontrollbeschreibung
    keeps its highest weight rather than being double-counted at both.

    Also uses IDF-style weighting (generic boilerplate words like "system",
    "prozess" count for little; rare/distinctive words count for much more)
    and gives partial credit for compound-word variants (e.g.
    "Kapazitätsmanagement" vs "Kapazitätsmonitoring" sharing a long prefix).
    """
    from .models import ReferenceReportEntry
    from collections import Counter
    import math

    ta_keywords = _translate_keywords(extract_keywords(test_activities_text))
    kb_keywords = _translate_keywords(extract_keywords(kontrollbeschreibung_text))

    # Merge into one keyword -> query-side weight map; a keyword found in
    # both fields keeps the higher (test_activities) weight.
    keyword_weights = {}
    for kw in kb_keywords:
        keyword_weights[kw] = max(keyword_weights.get(kw, 0), QUERY_FIELD_WEIGHT["kontrollbeschreibung"])
    for kw in ta_keywords:
        keyword_weights[kw] = max(keyword_weights.get(kw, 0), QUERY_FIELD_WEIGHT["test_activities"])

    if not keyword_weights:
        return []
    keywords = list(keyword_weights.keys())

    entries = list(ReferenceReportEntry.objects.all().select_related("reference_report"))
    if not entries:
        return []

    entry_field_data = []
    for entry in entries:
        tp_counts = Counter(_expand_with_compound_roots(WORD_RE.findall(entry.test_performed.lower())))
        kb_counts = Counter(_expand_with_compound_roots(WORD_RE.findall(entry.kontrollbeschreibung.lower())))
        entry_field_data.append((entry, tp_counts, kb_counts))

    # Document frequency across both fields combined, for IDF weighting
    n_total = len(entries)
    doc_freq = Counter()
    for kw in keywords:
        for _, tp_counts, kb_counts in entry_field_data:
            combined_words = set(tp_counts) | set(kb_counts)
            has_exact = kw in combined_words
            has_prefix = any(_best_partial_match_len(kw, w) >= 6 for w in combined_words)
            if has_exact or has_prefix:
                doc_freq[kw] += 1

    def idf(kw):
        # Standard BM25 IDF -- unlike a "+1"-smoothed version (which never lets
        # a term's weight drop below ~1 no matter how common it is), this
        # naturally suppresses terms appearing in nearly every entry down
        # towards zero. That matters a lot here: audit working papers are
        # heavily templated ("Gespräch mit dem Prozessverantwortlichen zur
        # Aufnahme des ... Prozesses. Einsichtnahme in die vorhandene
        # Prozessdokumentation...") -- without properly suppressing this
        # boilerplate, short generic entries could win purely by matching it,
        # especially once BM25 length-normalization is added (which inherently
        # favors short documents for any given match).
        df = doc_freq.get(kw, 0)
        raw = math.log((n_total - df + 0.5) / (df + 0.5) + 1)
        return max(raw, 0.05)  # tiny positive floor -- avoid exact zero/negative, not a meaningful weight either way

    # BM25-style length normalization -- WITHOUT this, a long/verbose entry
    # (lots of generic boilerplate audit language) accumulates many small
    # weak matches simply because it has more text, and can out-score a
    # SHORT entry that matches precisely on the actual topic. This is
    # exactly what real BM25 solves via document-length normalization;
    # dropped when this moved off FTS5, re-added here properly.
    K1, B = 1.5, 0.75
    tp_lengths = [sum(tp_counts.values()) for _, tp_counts, _ in entry_field_data]
    kb_lengths = [sum(kb_counts.values()) for _, _, kb_counts in entry_field_data]
    avg_tp_len = (sum(tp_lengths) / len(tp_lengths)) or 1
    avg_kb_len = (sum(kb_lengths) / len(kb_lengths)) or 1

    def field_score(kw, word_counts, doc_len, avg_len, weight):
        if kw in word_counts:
            tf = word_counts[kw]
            norm_tf = (tf * (K1 + 1)) / (tf + K1 * (1 - B + B * (doc_len / avg_len))) if doc_len else tf
            return norm_tf * weight * 6  # exact match, length-normalized
        best_match = 0
        for w in word_counts:
            best_match = max(best_match, _best_partial_match_len(kw, w))
        if best_match >= 6:
            # Reward absolute match length, not a ratio normalized by full word
            # length -- a 10-char match (prefix or embedded substring) is a
            # strong, distinctive signal regardless of the rest of each word's
            # length. Lightly length-normalized too, so a long doc's one
            # partial match doesn't get an unfair boost relative to a short doc's.
            length_penalty = avg_len / max(doc_len, avg_len)
            return weight * (best_match - 5) * (0.7 + 0.3 * length_penalty)
        return 0

    scored = []
    for entry, tp_counts, kb_counts in entry_field_data:
        tp_len = sum(tp_counts.values())
        kb_len = sum(kb_counts.values())
        score = 0.0
        for kw in keywords:
            base_weight = idf(kw) * keyword_weights[kw]
            score += field_score(kw, tp_counts, tp_len, avg_tp_len, base_weight * ENTRY_FIELD_WEIGHT["test_performed"])
            score += field_score(kw, kb_counts, kb_len, avg_kb_len, base_weight * ENTRY_FIELD_WEIGHT["kontrollbeschreibung"])
        if score > 0:
            scored.append((score, entry))

    scored.sort(key=lambda x: -x[0])
    return [_entry_to_suggestion(entry, round(score, 2)) for score, entry in scored[:top_n]]


# ---------------------------------------------------------------------------
# 8. Custom placeholder date formatting (per-placeholder, configurable EN/DE style)
# ---------------------------------------------------------------------------

GERMAN_MONTHS = ["Januar", "Februar", "März", "April", "Mai", "Juni", "Juli",
                 "August", "September", "Oktober", "November", "Dezember"]


def _ordinal_suffix(day):
    if 11 <= day % 100 <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def format_custom_date(date_obj, date_format):
    """
    Formats a date for a custom placeholder per its configured style:
      - en_ordinal: "8th October 2025"
      - de_ordinal: "1. Januar 2025"
    """
    if date_format == "de_ordinal":
        return f"{date_obj.day}. {GERMAN_MONTHS[date_obj.month - 1]} {date_obj.year}"
    day = date_obj.day
    return f"{day}{_ordinal_suffix(day)} {date_obj.strftime('%B')} {date_obj.year}"
